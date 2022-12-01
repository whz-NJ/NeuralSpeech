import os

from docx import Document
import win32com
import win32com.client
import preprocess

sports_corpus_root_dir = "../sports_corpus"
processed_sports_corpus_root_dir = "../sports_corpus2"
# corpus_root_dir = r"C:\Users\OS\Desktop\足球"
# processed_corpus_root_dir = r"C:\Users\OS\Desktop\足球2"
# aiui_football_dir = "../noised_aiui_football"
# processed_aiui_football_dir = "../aiui_football2"
# std_aiui_football_dir = "../std_noised_aiui_football2" #保存标准化后的正确语料-ASR语料对列表文件的目录
aiui_football_dir = "../noised_aiui_football" #包含 *_asr.txt 和 *.txt 两种文件
processed_aiui_football_dir = "../aiui_football2" # 只有标准化的 *.txt文件
std_aiui_football_dir = "../std_noised_aiui_football2" #保存标准化后的正确语料-ASR语料对列表文件的目录(std_*_asr.txt文件)
# aiui_football_dir = r'C:\corpus\2022_10\new' #包含 *_asr.txt 和 *.txt 两种文件
# processed_aiui_football_dir = r'C:\corpus\2022_10\new2' # 只有标准化的 *.txt文件
# std_aiui_football_dir = r'C:\corpus\2022_10\std_new2' #保存标准化后的正确语料-ASR语料对列表文件的目录(std_*_asr.txt文件)

def replace_dot_path(path):
    result = ""
    pwd = os.getcwd()
    if path.startswith(".."):
        ppwd = os.path.dirname(pwd)
        result = os.path.join(ppwd, path[3:]) #子目录不能包含最开始的 /
    elif path.startswith("."):
        result = os.path.join(pwd, path[2:]) #子目录不能包含最开始的 /
    else:
        result = path
    return result
sports_corpus_root_dir = replace_dot_path(sports_corpus_root_dir)
processed_sports_corpus_root_dir = replace_dot_path(processed_sports_corpus_root_dir)
aiui_football_dir = replace_dot_path(aiui_football_dir)
processed_aiui_football_dir = replace_dot_path(processed_aiui_football_dir)
std_aiui_football_dir = replace_dot_path(std_aiui_football_dir)

def open_docx_wps(path):
    try:
        doc = Document(path) #先用docx打开
    except: #打开失败，用wps打开
        try:
            wps = win32com.client.gencache.EnsureDispatch('kwps.application')
        except:
            wps = win32com.client.gencache.EnsureDispatch('wps.application')
        doc = wps.Documents.Open(path)
        pwd = os.getcwd()
        tmp_dir = os.path.join(pwd, 'tmp')
        if not os.path.exists(tmp_dir):
            os.mkdir(tmp_dir)
        new_path = os.path.join(tmp_dir, os.path.basename(path))
        doc.SaveAs2(new_path, 12) #另存为新的docx文件
        wps.Documents.Close(win32com.client.constants.wdDoNotSaveChanges)
        wps.Quit
        os.remove(path) #老的docx文件删除
        os.rename(new_path, path) #文件还原到源语料路径下
    doc = Document(path) #再次尝试用docx打开
    return doc

def loadSportsCorpus(root_dir):
    for root,dirs,files in os.walk(root_dir):
        for file in files:
            if root.find('五大联赛') == -1 and root.find('世界杯') == -1:  # 只处理足球相关语料
                continue
            if file.find('滑雪') != -1 or file.find('雪车') != -1: #只处理足球相关语料
                continue
            file_path = os.path.join(root, file)
            saved_corpus_path = file_path.replace(sports_corpus_root_dir, processed_sports_corpus_root_dir)
            saved_corpus_dir = os.path.dirname(saved_corpus_path)
            if not os.path.exists(saved_corpus_dir):
                os.makedirs(saved_corpus_dir)

            if saved_corpus_path.endswith(".docx"):
                saved_corpus_path = saved_corpus_path[:-4] + "txt"
            elif saved_corpus_path.endswith(".doc"):
                saved_corpus_path = saved_corpus_path[:-3] + "txt"
            # if os.path.exists(saved_corpus_path): #跳过已经处理过的文件
            #     continue

            corpus_list = []
            if file.endswith(".doc"):
                try:
                    wps = win32com.client.gencache.EnsureDispatch('kwps.application')
                except:
                    wps = win32com.client.gencache.EnsureDispatch('wps.application')
                print("opening doc file: " + file_path)
                doc = wps.Documents.Open(file_path)
                new_file_path = file_path[:-3] + "docx"
                doc.SaveAs2(new_file_path, 12)
                #doc.Close()
                #wps.Documents.Close()
                wps.Documents.Close(win32com.client.constants.wdDoNotSaveChanges)
                wps.Quit
                os.remove(file_path)
                file_path = new_file_path
            if file_path.endswith(".docx"):
                doc = open_docx_wps(file_path)
                for paragraph in doc.paragraphs[1:]:
                    sentences = preprocess.normAndTokenize(paragraph.text, min_sentence_len=2)
                    for sentence in sentences:
                        sentence = sentence.replace(" ", "")
                        corpus_list.append(sentence + "\n")
                #doc.Close()
            elif file_path.endswith(".txt"):
                with open(file_path, 'r', encoding='utf-8') as infile:
                    for line in infile.readlines():
                        sentences = preprocess.normAndTokenize(line, min_sentence_len=2)
                        for sentence in sentences:
                            sentence = sentence.replace(" ", "")
                            corpus_list.append(sentence + "\n")
            with open(saved_corpus_path, 'w', encoding='utf-8') as outfile:
                outfile.writelines(corpus_list)

def loadAiuiFootballCorpus(root_dir):
    for root,dirs,files in os.walk(root_dir):
        for file in files:
            file_path = os.path.join(root, file)
            if file.endswith("_asr.txt"): #有正确语料和ASR输出语料(但未标准化/分句)
                saved_corpus_path = file_path.replace(aiui_football_dir, std_aiui_football_dir)
                saved_corpus_dir = os.path.dirname(saved_corpus_path)
                # 原文件名前面加上 std_ 前缀
                saved_corpus_file_name = "std_" + os.path.basename(saved_corpus_path)
                saved_corpus_path = os.path.join(saved_corpus_dir, saved_corpus_file_name)
            else: #只有正确语料(但未标准化/分句)
                saved_corpus_path = file_path.replace(aiui_football_dir, processed_aiui_football_dir)
                saved_corpus_dir = os.path.dirname(saved_corpus_path)
            if not os.path.exists(saved_corpus_dir):
                os.makedirs(saved_corpus_dir)

            corpus_list = []
            with open(file_path, 'r', encoding='utf-8') as infile:
                    if file.endswith("_asr.txt"): #有正确语料和ASR输出语料，开始标准化
                        for line in infile.readlines():
                            fields = line.split("\t")
                            if len(fields) != 2:
                                continue
                            ref_sentences = fields[0]
                            hypo_sentences = fields[1]
                            ref_sentences = preprocess.normAndTokenize(ref_sentences, min_sentence_len=2)
                            hypo_sentences = preprocess.normAndTokenize(hypo_sentences, min_sentence_len=2)
                            if not ref_sentences or not hypo_sentences or len(ref_sentences) != len(hypo_sentences):
                                continue
                            for ref_sentence, hypo_sentence in zip(ref_sentences, hypo_sentences):
                                ref_hypo = ref_sentence + "\t" + hypo_sentence
                                corpus_list.append(ref_hypo + "\n")
                    else: #只有正确语料，开始标准化
                        for line in infile.readlines():
                            corpus = line.strip()
                            sentences = preprocess.normAndTokenize(corpus, min_sentence_len=2)
                            for sentence in sentences:
                                sentence = sentence.replace(" ", "") #删除中间的空格
                                corpus_list.append(sentence + "\n")
            with open(saved_corpus_path, 'w', encoding='utf-8') as outfile:
                outfile.writelines(corpus_list)

loadSportsCorpus(sports_corpus_root_dir)
loadAiuiFootballCorpus(aiui_football_dir)
